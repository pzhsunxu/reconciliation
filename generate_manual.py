# -*- coding: utf-8 -*-
"""生成对账系统使用手册 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_table_with_header(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def add_step(doc, number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'步骤{number}: ')
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(text).font.size = Pt(11)


def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('提示: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(10)
    p.add_run(text).font.size = Pt(10)


def generate_manual():
    doc = Document()

    # Title
    title = doc.add_heading('对账系统使用手册', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('版本: 1.0').font.size = Pt(11)
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('日期: 2026-05-16').font.size = Pt(11)
    doc.add_paragraph('')

    # ========== Chapter 1 ==========
    h = doc.add_heading('1. 系统简介', level=1)
    doc.add_paragraph(
        '对账系统是公司用于管理全国酒店各平台销售数据的综合管理平台。'
        '系统支持自动对账和手动对账，可自动生成对账报表并支持人工复核签字。'
        '主要功能包括酒店管理、平台账号管理、销售数据查询、日常开支管理、对账任务和报表中心。'
    )

    # ========== Chapter 2 ==========
    h = doc.add_heading('2. 系统启动', level=1)

    h2 = doc.add_heading('2.1 环境要求', level=2)
    doc.add_paragraph('Python 3.10 及以上版本')
    doc.add_paragraph('Node.js 16 及以上版本')

    h2 = doc.add_heading('2.2 启动步骤', level=2)
    add_step(doc, 1, '打开终端，进入项目目录：cd reconciliation')
    add_step(doc, 2, '安装后端依赖：cd backend && pip install -r requirements.txt')
    add_step(doc, 3, '启动后端服务：python main.py （默认端口8001）')
    add_step(doc, 4, '打开新终端，安装前端依赖：cd frontend && npm install')
    add_step(doc, 5, '启动前端服务：npm run dev （默认端口5173）')
    add_step(doc, 6, '浏览器访问 http://localhost:5173 即可使用系统')

    add_tip(doc, '后端API文档地址：http://localhost:8001/docs，可查看所有接口的详细说明。')

    # ========== Chapter 3 ==========
    h = doc.add_heading('3. 首页仪表盘', level=1)

    doc.add_paragraph(
        '进入系统后默认显示首页仪表盘，展示以下核心指标：'
    )

    dash_headers = ['指标', '说明']
    dash_rows = [
        ['酒店总数', '系统中所有酒店的数量'],
        ['分润酒店', '采用分润模式的酒店数量'],
        ['对账任务数', '已创建的对账任务总数'],
        ['待复核报表', '尚未复核的报表数量'],
        ['本月销售额', '当前月份所有酒店的总销售额'],
        ['各平台销售额', '美团/携程/飞猪/抖音/PMS各自的销售额'],
    ]
    add_table_with_header(doc, dash_headers, dash_rows)
    doc.add_paragraph('')

    # ========== Chapter 4 ==========
    h = doc.add_heading('4. 酒店管理', level=1)

    doc.add_paragraph('路径：左侧菜单 → 酒店管理')

    h2 = doc.add_heading('4.1 添加酒店', level=2)
    add_step(doc, 1, '点击"新增"按钮，打开新增酒店弹窗')
    add_step(doc, 2, '填写酒店名称（必填）')
    add_step(doc, 3, '填写酒店地址（选填）')
    add_step(doc, 4, '选择合作模式：分润 或 全租')
    add_step(doc, 5, '若选择"分润"模式，设置公司占比和房东占比（两者之和应为1.0）')
    add_step(doc, 6, '选择启用/停用状态')
    add_step(doc, 7, '点击"确定"完成创建')

    add_tip(doc, '分润模式下，公司占比+房东占比必须等于100%。全租模式下公司占比固定为100%。')

    h2 = doc.add_heading('4.2 查询酒店', level=2)
    doc.add_paragraph('在查询条件中输入酒店名称或选择合作模式，点击"查询"即可筛选。')

    h2 = doc.add_heading('4.3 编辑酒店', level=2)
    doc.add_paragraph('在酒店列表中找到目标酒店，点击"编辑"按钮，修改后点击"确定"保存。')

    h2 = doc.add_heading('4.4 删除酒店', level=2)
    doc.add_paragraph('点击"删除"按钮，确认后删除酒店。')
    add_tip(doc, '删除酒店将同时关联的平台账号一并删除，请谨慎操作。')

    # ========== Chapter 5 ==========
    h = doc.add_heading('5. 平台账号管理', level=1)

    doc.add_paragraph('路径：左侧菜单 → 平台账号')

    h2 = doc.add_heading('5.1 添加平台账号', level=2)
    add_step(doc, 1, '点击"新增"按钮')
    add_step(doc, 2, '选择所属酒店')
    add_step(doc, 3, '选择平台类型（美团/携程/飞猪/抖音/PMS）')
    add_step(doc, 4, '填写账号名称')
    add_step(doc, 5, '填写平台账号ID（选填）')
    add_step(doc, 6, '设置启用状态')
    add_step(doc, 7, '点击"确定"完成创建')

    h2 = doc.add_heading('5.2 查询与编辑', level=2)
    doc.add_paragraph('可通过酒店筛选平台账号列表，点击"编辑"修改，点击"删除"移除账号。')

    # ========== Chapter 6 ==========
    h = doc.add_heading('6. 销售数据', level=1)

    doc.add_paragraph('路径：左侧菜单 → 销售数据')

    h2 = doc.add_heading('6.1 查询销售数据', level=2)
    add_step(doc, 1, '在查询条件中选择酒店、平台、起始日期、结束日期')
    add_step(doc, 2, '点击"查询"按钮')
    doc.add_paragraph('查询结果展示订单号、平台、入住/离店日期、房型、金额、佣金、净收入等信息。')

    h2 = doc.add_heading('6.2 拉取销售数据', level=2)
    add_step(doc, 1, '点击"拉取数据"按钮')
    add_step(doc, 2, '选择目标酒店')
    add_step(doc, 3, '设置起始日期和结束日期')
    add_step(doc, 4, '点击"拉取"，系统将从该酒店绑定的各平台拉取数据')
    add_tip(doc, '当前系统使用模拟数据生成销售记录，后续对接真实API后将自动拉取实际数据。')

    # ========== Chapter 7 ==========
    h = doc.add_heading('7. 日常开支管理', level=1)

    doc.add_paragraph('路径：左侧菜单 → 日常开支')

    h2 = doc.add_heading('7.1 录入开支', level=2)
    add_step(doc, 1, '点击"新增"按钮')
    add_step(doc, 2, '选择所属酒店')
    add_step(doc, 3, '选择开支类别：水电 / 人工 / 维修 / 清洁 / 其他')
    add_step(doc, 4, '输入金额')
    add_step(doc, 5, '选择开支日期')
    add_step(doc, 6, '填写说明（选填）')
    add_step(doc, 7, '填写录入人姓名')
    add_step(doc, 8, '点击"确定"保存')

    h2 = doc.add_heading('7.2 查询与编辑', level=2)
    doc.add_paragraph(
        '可按酒店筛选开支列表。点击"编辑"修改记录，点击"删除"移除记录。'
        '日常开支将计入对账计算中的总开支项，影响分润金额。'
    )

    # ========== Chapter 8 ==========
    h = doc.add_heading('8. 对账管理', level=1)

    doc.add_paragraph('路径：左侧菜单 → 对账管理')

    h2 = doc.add_heading('8.1 自动对账', level=2)
    doc.add_paragraph(
        '系统默认每月12号和28号凌晨2:00自动执行对账任务，无需人工操作。'
        '自动对账流程：\n'
        '1. 系统计算对账日期区间\n'
        '2. 遍历所有分润模式酒店\n'
        '3. 从各平台拉取销售数据\n'
        '4. 计算每家酒店的分润金额\n'
        '5. 生成对账报表（状态为"草稿"）'
    )

    h2 = doc.add_heading('8.2 手动对账', level=2)
    add_step(doc, 1, '点击"手动对账"按钮')
    add_step(doc, 2, '设置对账起始日期和结束日期')
    add_step(doc, 3, '点击"确定"创建任务')
    add_step(doc, 4, '任务创建后在后台自动执行，可在列表中查看进度')
    doc.add_paragraph('任务状态说明：')
    status_headers = ['状态', '说明']
    status_rows = [
        ['待执行', '任务已创建，等待执行'],
        ['执行中', '正在拉取数据并计算'],
        ['已完成', '对账计算完成，报表已生成'],
        ['失败', '执行出错，可查看错误信息'],
    ]
    add_table_with_header(doc, status_headers, status_rows)
    doc.add_paragraph('')

    # ========== Chapter 9 ==========
    h = doc.add_heading('9. 报表中心', level=1)

    doc.add_paragraph('路径：左侧菜单 → 报表中心')

    h2 = doc.add_heading('9.1 查看报表', level=2)
    add_step(doc, 1, '在报表列表中可选择酒店、状态进行筛选')
    add_step(doc, 2, '点击"详情"按钮查看报表完整信息')
    doc.add_paragraph('报表详情包含以下信息：')
    report_fields = [
        '酒店名称', '对账周期', '总销售额', '平台佣金',
        '净收入', '日常开支', '公司应得', '房东应得',
        '复核人', '复核时间',
    ]
    for f in report_fields:
        doc.add_paragraph(f, style='List Bullet')

    h2 = doc.add_heading('9.2 复核签字', level=2)
    add_step(doc, 1, '在报表列表中找到待复核的报表（状态为"草稿"）')
    add_step(doc, 2, '点击"复核"按钮')
    add_step(doc, 3, '输入复核人姓名')
    add_step(doc, 4, '点击"确定"完成复核')
    doc.add_paragraph('复核后报表状态变为"已复核"，不可再次修改。')
    add_tip(doc, '复核人姓名会记录在报表中，请确保输入正确。')

    h2 = doc.add_heading('9.3 导出报表', level=2)
    add_step(doc, 1, '在报表列表中点击"导出"按钮')
    add_step(doc, 2, '系统将自动生成Excel文件并下载')
    doc.add_paragraph('Excel报表包含完整的对账信息和复核记录。')

    # ========== Chapter 10 ==========
    h = doc.add_heading('10. 常见问题', level=1)

    faq_headers = ['问题', '解决方案']
    faq_rows = [
        ['对账任务显示"失败"', '检查错误信息，确认酒店已绑定平台账号且处于启用状态'],
        ['报表金额为0', '确认对账周期内有销售数据，可先手动拉取数据再执行对账'],
        ['分润比例如何设置', '在酒店管理中设置公司占比和房东占比，两者之和必须等于1.0'],
        ['全租酒店为什么没有报表', '全租模式不涉及分润，系统仅对分润模式酒店生成报表'],
        ['如何修改已复核的报表', '已复核报表不可修改，需联系管理员创建新的对账任务'],
        ['前端页面打不开', '确认后端服务是否正常启动（http://localhost:8001/docs可访问）'],
    ]
    add_table_with_header(doc, faq_headers, faq_rows)

    # ========== Chapter 11 ==========
    h = doc.add_heading('11. 对账计算说明', level=1)

    doc.add_paragraph(
        '对账计算的核心公式如下：\n\n'
        '① 总销售额 = 该对账周期内所有平台订单的销售金额合计\n\n'
        '② 总佣金 = 该对账周期内所有平台订单的平台佣金合计\n\n'
        '③ 净收入 = 总销售额 - 总佣金\n\n'
        '④ 总开支 = 该对账周期内该酒店的所有日常开支合计\n\n'
        '⑤ 可分配利润 = 净收入 - 总开支\n\n'
        '⑥ 公司应得 = 可分配利润 × 公司占比\n\n'
        '⑦ 房东应得 = 可分配利润 × 房东占比\n\n'
        '验证：公司应得 + 房东应得 = 可分配利润'
    )

    doc.add_paragraph('')
    h2 = doc.add_heading('11.1 计算示例', level=2)
    doc.add_paragraph(
        '假设某酒店某周期内：\n'
        '总销售额 = 100,000 元\n'
        '总佣金 = 10,000 元\n'
        '净收入 = 100,000 - 10,000 = 90,000 元\n'
        '总开支 = 5,000 元（水电3000 + 人工2000）\n'
        '可分配利润 = 90,000 - 5,000 = 85,000 元\n'
        '公司应得 = 85,000 × 60% = 51,000 元\n'
        '房东应得 = 85,000 × 40% = 34,000 元\n'
        '验证：51,000 + 34,000 = 85,000 元 ✓'
    )

    doc.save('对账系统使用手册.docx')
    print('User manual saved: 对账系统使用手册.docx')


if __name__ == '__main__':
    generate_manual()
