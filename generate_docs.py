# -*- coding: utf-8 -*-
"""生成对账系统设计方案 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def set_table_style(doc, table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_heading(doc, title):
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_sub_heading(doc, title):
    h = doc.add_heading(title, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table_with_header(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)
    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    set_table_style(doc, table)
    return table


def generate_design_doc():
    doc = Document()

    # Title
    title = doc.add_heading('对账系统设计方案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('版本: 1.1').font.size = Pt(11)
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('日期: 2026-05-17').font.size = Pt(11)
    doc.add_paragraph('')

    # ========== Chapter 1 ==========
    add_section_heading(doc, '1. 项目概述')

    add_sub_heading(doc, '1.1 项目背景')
    doc.add_paragraph(
        '公司在全国运营一百多家酒店，每个酒店在美团、携程、飞猪、抖音、PMS等平台上拥有独立账号。'
        '为规范管理销售费用、明确公司与房东之间的分润比例，需建设一套对账系统来实现对各平台销售数据的采集、'
        '计算与报表生成，支持人工复核签字流程。'
    )

    add_sub_heading(doc, '1.2 业务场景')
    doc.add_paragraph(
        '酒店合作模式分为两种：\n'
        '（1）分润模式：酒店房东与公司合作，房东占比40%，公司占比60%（比例可配置），需对账分润。\n'
        '（2）全租模式：公司100%承租酒店，不涉及对账。'
    )

    add_sub_heading(doc, '1.3 系统目标')
    goals = [
        '统一管理全国酒店的各平台销售数据',
        '支持自动对账（每月12号/28号）和手动对账',
        '支持日常开支的人工录入管理',
        '自动生成对账报表，支持复核签字和Excel导出',
        '为后续对接真实平台API预留扩展接口',
    ]
    for g in goals:
        doc.add_paragraph(g, style='List Bullet')

    # ========== Chapter 2 ==========
    add_section_heading(doc, '2. 技术架构')

    add_sub_heading(doc, '2.1 技术栈')
    tech_headers = ['层次', '技术选型', '版本']
    tech_rows = [
        ['后端框架', 'FastAPI', '0.124+'],
        ['数据库', 'SQLite (SQLAlchemy ORM)', '2.0+'],
        ['定时任务', 'APScheduler', '3.11+'],
        ['Excel导出', 'openpyxl', '3.1+'],
        ['前端框架', 'Vue 3', '3.4+'],
        ['UI组件库', 'Element Plus', '2.8+'],
        ['构建工具', 'Vite', '5.4+'],
        ['测试框架', 'pytest', '9.0+'],
    ]
    add_table_with_header(doc, tech_headers, tech_rows)
    doc.add_paragraph('')

    add_sub_heading(doc, '2.2 系统架构图')
    doc.add_paragraph(
        '系统采用前后端分离架构：\n\n'
        '前端层（Vue 3 + Element Plus）：负责页面展示、用户交互、API请求\n'
        '    ↓ HTTP / JSON\n'
        '后端层（FastAPI）：RESTful API接口、业务逻辑处理、数据校验\n'
        '    ↓ SQLAlchemy ORM\n'
        '数据层（SQLite）：持久化存储所有业务数据\n\n'
        '定时任务层（APScheduler）：独立于API请求，每月触发自动对账流程'
    )

    add_sub_heading(doc, '2.3 项目目录结构')
    doc.add_paragraph(
        'reconciliation/\n'
        '├── backend/\n'
        '│   ├── app/\n'
        '│   │   ├── main.py              # FastAPI应用入口\n'
        '│   │   ├── config.py            # 配置管理\n'
        '│   │   ├── database.py          # 数据库连接\n'
        '│   │   ├── models.py            # SQLAlchemy数据模型\n'
        '│   │   ├── schemas.py           # Pydantic请求/响应模型\n'
        '│   │   ├── api/                 # RESTful API路由\n'
        '│   │   │   ├── hotels.py        # 酒店管理接口\n'
        '│   │   │   ├── platforms.py     # 平台账号管理接口\n'
        '│   │   │   ├── sales.py         # 销售数据管理接口\n'
        '│   │   │   ├── expenses.py      # 日常开支管理接口\n'
        '│   │   │   ├── reconciliation.py # 对账任务接口\n'
        '│   │   │   └── reports.py       # 对账报表接口\n'
        '│   │   ├── services/            # 业务服务层\n'
        '│   │   │   ├── platform_service.py\n'
        '│   │   │   ├── reconciliation_service.py\n'
        '│   │   │   ├── report_service.py\n'
        '│   │   │   └── scheduler.py\n'
        '│   │   └── utils/\n'
        '│   │       └── excel_export.py\n'
        '│   ├── requirements.txt\n'
        '│   └── main.py                  # 启动入口\n'
        '├── frontend/                     # 前端项目\n'
        '├── tests/                        # 测试用例\n'
        '└── 对账系统需求.md\n',
        style='No Spacing'
    )

    # ========== Chapter 3 ==========
    add_section_heading(doc, '3. 数据模型设计')

    add_sub_heading(doc, '3.1 Hotel - 酒店表')
    doc.add_paragraph('存储酒店基本信息及合作模式配置。')
    hotel_headers = ['字段', '类型', '说明']
    hotel_rows = [
        ['id', 'Integer', '主键'],
        ['name', 'String(100)', '酒店名称'],
        ['location', 'String(200)', '酒店地址'],
        ['cooperation_type', 'Enum', 'split(分租) / full(全租)'],
        ['company_share', 'Float', '公司占比(如0.6)'],
        ['owner_share', 'Float', '房东占比(如0.4)'],
        ['status', 'Boolean', '启用/停用'],
        ['created_at', 'DateTime', '创建时间'],
    ]
    add_table_with_header(doc, hotel_headers, hotel_rows)
    doc.add_paragraph('')

    add_sub_heading(doc, '3.2 PlatformAccount - 平台账号表')
    doc.add_paragraph('存储酒店在各平台的账号信息。')
    plat_headers = ['字段', '类型', '说明']
    plat_rows = [
        ['id', 'Integer', '主键'],
        ['hotel_id', 'ForeignKey', '关联酒店'],
        ['platform', 'Enum', 'meituan/ctrip/fliggy/douyin/pms'],
        ['account_name', 'String(100)', '账号名称'],
        ['account_id', 'String(100)', '平台账号ID'],
        ['api_config', 'JSON', 'API配置信息(预留)'],
        ['status', 'Boolean', '启用/停用'],
    ]
    add_table_with_header(doc, plat_headers, plat_rows)
    doc.add_paragraph('')

    add_sub_heading(doc, '3.3 SalesData - 销售数据表')
    doc.add_paragraph('存储从各平台拉取的销售订单数据。')
    sales_headers = ['字段', '类型', '说明']
    sales_rows = [
        ['id', 'Integer', '主键'],
        ['hotel_id', 'ForeignKey', '关联酒店'],
        ['platform', 'Enum', '数据来源平台'],
        ['order_no', 'String(100)', '订单号'],
        ['check_in', 'Date', '入住日期'],
        ['check_out', 'Date', '离店日期'],
        ['room_no', 'String(50)', '房号'],
        ['room_type', 'String(50)', '房型'],
        ['amount', 'Decimal', '销售金额'],
        ['commission', 'Decimal', '平台佣金'],
        ['net_amount', 'Decimal', '净收入'],
        ['source', 'Enum', 'auto(自动) / manual(手动)'],
        ['pulled_at', 'DateTime', '拉取时间'],
    ]
    add_table_with_header(doc, sales_headers, sales_rows)
    doc.add_paragraph('')

    add_sub_heading(doc, '3.4 ExpenseItem - 日常开支表')
    exp_headers = ['字段', '类型', '说明']
    exp_rows = [
        ['id', 'Integer', '主键'],
        ['hotel_id', 'ForeignKey', '关联酒店'],
        ['category', 'Enum', 'utility/labor/maintenance/cleaning/other'],
        ['amount', 'Decimal', '金额'],
        ['description', 'String(500)', '说明'],
        ['expense_date', 'Date', '开支日期'],
        ['created_by', 'String(50)', '录入人'],
        ['created_at', 'DateTime', '创建时间'],
    ]
    add_table_with_header(doc, exp_headers, exp_rows)
    doc.add_paragraph('')

    add_sub_heading(doc, '3.5 ReconciliationJob - 对账任务表')
    job_headers = ['字段', '类型', '说明']
    job_rows = [
        ['id', 'Integer', '主键'],
        ['job_type', 'Enum', 'auto(自动) / manual(手动)'],
        ['start_date', 'Date', '对账起始日期'],
        ['end_date', 'Date', '对账结束日期'],
        ['status', 'Enum', 'pending/running/completed/failed'],
        ['hotels_count', 'Integer', '涉及酒店数量'],
        ['created_at', 'DateTime', '创建时间'],
        ['completed_at', 'DateTime', '完成时间'],
        ['error_msg', 'Text', '错误信息'],
    ]
    add_table_with_header(doc, job_headers, job_rows)
    doc.add_paragraph('')

    add_sub_heading(doc, '3.6 ReconciliationReport - 对账报表表')
    rep_headers = ['字段', '类型', '说明']
    rep_rows = [
        ['id', 'Integer', '主键'],
        ['job_id', 'ForeignKey', '关联对账任务'],
        ['hotel_id', 'ForeignKey', '关联酒店'],
        ['period_start', 'Date', '周期起始'],
        ['period_end', 'Date', '周期结束'],
        ['total_sales', 'Decimal', '总销售额'],
        ['total_commission', 'Decimal', '总平台佣金'],
        ['total_net_income', 'Decimal', '总净收入'],
        ['total_expense', 'Decimal', '总日常开支'],
        ['company_amount', 'Decimal', '公司应得分润'],
        ['owner_amount', 'Decimal', '房东应得分润'],
        ['status', 'Enum', 'draft(草稿) / reviewed(已复核)'],
        ['reviewer', 'String(50)', '复核人'],
        ['reviewed_at', 'DateTime', '复核时间'],
        ['created_at', 'DateTime', '创建时间'],
    ]
    add_table_with_header(doc, rep_headers, rep_rows)

    # ========== Chapter 4 ==========
    add_section_heading(doc, '4. API接口设计')

    api_headers = ['方法', '路径', '说明']
    api_rows = [
        ['GET', '/api/hotels', '酒店列表(分页+筛选)'],
        ['POST', '/api/hotels', '新增酒店'],
        ['PUT', '/api/hotels/{id}', '修改酒店'],
        ['DELETE', '/api/hotels/{id}', '删除酒店'],
        ['GET', '/api/platforms', '平台账号列表'],
        ['POST', '/api/platforms', '新增平台账号'],
        ['PUT', '/api/platforms/{id}', '修改平台账号'],
        ['GET', '/api/sales', '销售数据列表'],
        ['POST', '/api/sales/pull', '手动拉取销售数据'],
        ['GET', '/api/expenses', '日常开支列表'],
        ['POST', '/api/expenses', '新增日常开支'],
        ['PUT', '/api/expenses/{id}', '修改开支'],
        ['DELETE', '/api/expenses/{id}', '删除开支'],
        ['GET', '/api/reconciliations', '对账任务列表'],
        ['POST', '/api/reconciliations', '创建手动对账任务'],
        ['GET', '/api/reconciliations/{id}', '对账任务详情'],
        ['GET', '/api/reports', '报表列表'],
        ['GET', '/api/reports/{id}', '报表详情'],
        ['POST', '/api/reports/{id}/review', '报表复核签字'],
        ['GET', '/api/reports/{id}/excel', '报表Excel导出'],
        ['GET', '/api/dashboard/stats', '仪表盘统计数据'],
    ]
    add_table_with_header(doc, api_headers, api_rows)

    # ========== Chapter 5 ==========
    add_section_heading(doc, '5. 核心业务逻辑')

    add_sub_heading(doc, '5.1 对账计算公式')
    doc.add_paragraph(
        '总销售额 = 该周期内所有平台订单的销售金额合计\n'
        '总佣金 = 该周期内所有平台订单的佣金合计\n'
        '净收入 = 总销售额 - 总佣金\n'
        '总开支 = 该周期内该酒店的所有日常开支合计\n'
        '可分配利润 = 净收入 - 总开支\n'
        '公司应得 = 可分配利润 × 公司占比\n'
        '房东应得 = 可分配利润 × 房东占比'
    )

    add_sub_heading(doc, '5.2 自动对账定时任务')
    doc.add_paragraph(
        '触发时间：每月12号和28号凌晨2:00\n\n'
        '执行流程：\n'
        '1. 计算对账日期区间（上一次对账结束日+1天至当天）\n'
        '2. 创建自动对账任务\n'
        '3. 遍历所有分润模式且启用的酒店\n'
        '4. 对每家酒店依次从各平台拉取销售数据\n'
        '5. 计算对账金额并生成报表\n'
        '6. 报表状态设为"草稿"待复核'
    )

    add_sub_heading(doc, '5.3 报表复核流程')
    doc.add_paragraph(
        '1. 对账任务完成后，报表状态为"草稿"\n'
        '2. 财务人员在报表中心查看报表详情\n'
        '3. 确认数据无误后，输入复核人姓名进行签字\n'
        '4. 报表状态变为"已复核"，记录复核人和时间\n'
        '5. 复核后的报表支持Excel导出'
    )

    # ========== Chapter 6 ==========
    add_section_heading(doc, '6. 前端页面设计')

    page_headers = ['页面', '路由', '功能说明']
    page_rows = [
        ['首页仪表盘', '/', '展示酒店总数、分润酒店数、对账任务数、待复核报表数、本月销售额及各平台销售额汇总'],
        ['酒店管理', '/hotels', '酒店的增删改查，设置合作模式和分润比例，启用/停用'],
        ['平台账号', '/platforms', '为酒店绑定各平台账号，管理账号信息'],
        ['销售数据', '/sales', '按条件查看/筛选各平台销售数据，支持手动拉取'],
        ['日常开支', '/expenses', '录入/编辑/删除酒店日常开支，按类别和日期筛选'],
        ['对账管理', '/reconciliation', '查看自动/手动对账任务进度，创建手动对账任务'],
        ['报表中心', '/reports', '查看对账报表详情，执行复核签字，导出Excel'],
    ]
    add_table_with_header(doc, page_headers, page_rows)

    # ========== Chapter 7 ==========
    add_section_heading(doc, '7. 部署方案')
    doc.add_paragraph(
        '开发环境：\n'
        '  后端：cd backend && python main.py (端口8001)\n'
        '  前端：cd frontend && npm run dev (端口5173)\n\n'
        '生产环境建议：\n'
        '  后端：gunicorn uvicorn.workers.UvicornWorker:app.main:app -w 4 -b 0.0.0.0:8000\n'
        '  前端：npm run build 后通过Nginx提供静态文件服务\n'
        '  数据库：将SQLite替换为PostgreSQL或MySQL\n'
        '  域名：配置反向代理，统一入口'
    )

    # ========== Chapter 8 ==========
    add_section_heading(doc, '8. 测试情况')
    add_sub_heading(doc, '8.1 单元测试')
    doc.add_paragraph(
        '测试框架：pytest 9.0\n'
        '测试用例总数：36\n'
        '通过数：36\n'
        '通过率：100%'
    )
    test_headers = ['测试模块', '用例数', '通过率']
    test_rows = [
        ['test_dashboard.py', '2', '100%'],
        ['test_hotel.py', '8', '100%'],
        ['test_platform.py', '5', '100%'],
        ['test_sales.py', '4', '100%'],
        ['test_expense.py', '5', '100%'],
        ['test_reconciliation.py', '5', '100%'],
        ['test_report.py', '7', '100%'],
        ['合计', '36', '100%'],
    ]
    add_table_with_header(doc, test_headers, test_rows)

    add_sub_heading(doc, '8.2 系统测试')
    doc.add_paragraph(
        '已完成端到端集成测试，验证了以下完整流程：\n'
        '1. 创建酒店（分润/全租模式）→ 成功\n'
        '2. 绑定5个平台账号 → 成功\n'
        '3. 录入日常开支 → 成功\n'
        '4. 仪表盘统计展示 → 数据正确\n'
        '5. 手动对账执行 → 任务完成，报表生成\n'
        '6. 报表复核签字 → 状态变更正确\n'
        '7. 前端页面功能 → 全部可用'
    )

    # ========== Chapter 9 ==========
    add_section_heading(doc, '9. 后续扩展')
    extensions = [
        '对接美团/携程/飞猪/抖音/PMS真实API，替换模拟数据',
        '添加用户登录和角色权限管理（管理员/操作员/复核员）',
        '支持CSV/Excel批量导入销售数据',
        '报表PDF导出和邮件发送',
        '对账数据趋势分析和可视化图表',
        '数据库升级至PostgreSQL/MySQL支持多用户并发',
    ]
    for e in extensions:
        doc.add_paragraph(e, style='List Bullet')

    doc.save('对账系统设计方案.docx')
    print('Design document saved: 对账系统设计方案.docx')


if __name__ == '__main__':
    generate_design_doc()
